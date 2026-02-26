"""File parsers — extract chapters from epub, pdf, docx, txt, md, mobi."""

import re, tempfile, warnings
from pathlib import Path
from bs4 import BeautifulSoup

warnings.filterwarnings("ignore", category=UserWarning, module="ebooklib")

# ── Chapter heading pattern ─────────────────────────────
_CH_RE = re.compile(
    r"^(?:"
    r"(?:chapter|ch\.?)\s+[\d\w]+[:\.\s\-—]*.*|"
    r"(?:part|section)\s+[\d\w]+[:\.\s\-—]*.*|"
    r"#{1,3}\s+.+"
    r")$",
    re.IGNORECASE | re.MULTILINE,
)


def parse_file(file_obj, filename: str) -> list[dict]:
    """Route to the correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    router = {
        ".txt": _parse_text, ".md": _parse_text,
        ".pdf": _parse_pdf, ".docx": _parse_docx,
        ".epub": _parse_epub, ".mobi": _parse_mobi,
    }
    fn = router.get(ext)
    if not fn:
        raise ValueError(f"Unsupported format: {ext}")
    return fn(file_obj)


def parse_pasted_text(text: str) -> list[dict]:
    return _detect_chapters(text)


# ── Shared helpers ──────────────────────────────────────

def _detect_chapters(text: str) -> list[dict]:
    """Split text on chapter headings; falls back to single chapter."""
    text = text.strip()
    markers = list(_CH_RE.finditer(text))
    if not markers:
        return [{"title": "Full Text", "text": text}]

    chapters = []
    # content before the first heading
    pre = text[: markers[0].start()].strip()
    if pre and len(pre) > 100:
        chapters.append({"title": "Preface", "text": pre})

    for i, m in enumerate(markers):
        end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        body = text[m.end() : end].strip()
        if body:
            chapters.append({
                "title": m.group().strip().lstrip("#").strip(),
                "text": body,
            })
    return chapters or [{"title": "Full Text", "text": text}]


def _clean_html(html_bytes) -> str:
    """Strip HTML to plain text."""
    if isinstance(html_bytes, str):
        html_bytes = html_bytes.encode()
    soup = BeautifulSoup(html_bytes, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _heading_from_html(html_bytes) -> str | None:
    soup = BeautifulSoup(html_bytes, "html.parser")
    h = soup.find(["h1", "h2", "h3"])
    return h.get_text(strip=True) if h else None


# ── Format parsers ──────────────────────────────────────

def _parse_text(f) -> list[dict]:
    raw = f.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="ignore")
    return _detect_chapters(raw)


def _parse_pdf(f) -> list[dict]:
    import pdfplumber
    pages = []
    with pdfplumber.open(f) as pdf:
        for p in pdf.pages:
            t = p.extract_text()
            if t:
                pages.append(t)
    return _detect_chapters("\n\n".join(pages))


def _parse_docx(f) -> list[dict]:
    from docx import Document
    doc = Document(f)
    chapters, title, body = [], None, []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            if body:
                chapters.append({"title": title or "Untitled", "text": "\n".join(body).strip()})
            title = para.text.strip()
            body = []
        elif para.text.strip():
            body.append(para.text)
    if body:
        chapters.append({"title": title or "Full Text", "text": "\n".join(body).strip()})
    if not chapters:
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return _detect_chapters(text)
    return chapters


def _parse_epub(f) -> list[dict]:
    import ebooklib
    from ebooklib import epub

    with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
        tmp.write(f.read() if hasattr(f, "read") else f)
        path = tmp.name
    try:
        book = epub.read_epub(path)

        # build ToC title lookup (href → title)
        toc_map = {}
        for entry in book.toc:
            if isinstance(entry, epub.Link):
                toc_map[entry.href.split("#")[0]] = entry.title
            elif isinstance(entry, tuple):
                for link in entry[1]:
                    if isinstance(link, epub.Link):
                        toc_map[link.href.split("#")[0]] = link.title

        chapters = []
        for item_id, _ in book.spine:
            item = book.get_item_with_id(item_id)
            if not item or item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue
            html = item.get_content()
            text = _clean_html(html)
            if len(text) < 50:
                continue
            fname = item.get_name()
            title = toc_map.get(fname) or _heading_from_html(html) or Path(fname).stem
            chapters.append({"title": title, "text": text})

        return chapters or [{"title": "Full Text", "text": "No readable content found."}]
    finally:
        Path(path).unlink(missing_ok=True)


def _parse_mobi(f) -> list[dict]:
    try:
        import mobi
    except ImportError:
        raise ImportError("Install the 'mobi' package, or convert .mobi → .epub via Calibre.")

    with tempfile.NamedTemporaryFile(suffix=".mobi", delete=False) as tmp:
        tmp.write(f.read() if hasattr(f, "read") else f)
        path = tmp.name
    try:
        tempdir, extracted = mobi.extract(path)
        if Path(extracted).suffix.lower() == ".epub":
            with open(extracted, "rb") as ef:
                return _parse_epub(ef)
        # older mobi → extracted HTML
        with open(extracted, "r", encoding="utf-8", errors="ignore") as hf:
            return _detect_chapters(_clean_html(hf.read().encode()))
    except Exception as e:
        raise RuntimeError(f"Mobi extraction failed: {e}. Convert to .epub with Calibre.") from e
    finally:
        Path(path).unlink(missing_ok=True)
